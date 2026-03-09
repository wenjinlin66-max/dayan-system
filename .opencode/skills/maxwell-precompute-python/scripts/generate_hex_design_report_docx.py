from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path("results/agv_pdf_run")
PRE = ROOT / "precompute_outputs"


def add_table_from_df(doc: Document, df: pd.DataFrame, title: str, max_rows: int = 20) -> None:
    doc.add_paragraph(title)
    if df.empty:
        doc.add_paragraph("(无数据)")
        return
    show = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(show.columns):
        hdr[i].text = str(c)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(show.columns):
            cells[i].text = str(row[c])


def main() -> None:
    out_docx = ROOT / "AGV六边形阵列式耦合机构参数设计报告.docx"

    fixed_df = pd.DataFrame(
        [
            ("工作频率 f", "85 kHz"),
            ("额定功率 P", "1~2 kW"),
            ("输入电流 Iin", "20 A"),
            ("传输距离 Air Gap", "50 mm"),
            ("线圈形状", "矩形线圈（后续用于阵列耦合优化）"),
            ("阵列数量", "3（1x3）"),
            ("线材类型", "绞线 Stranded"),
            ("负载电阻 Rload", "10 ohm"),
        ],
        columns=["参数", "取值"],
    )

    scan_100_200 = pd.read_csv(PRE / "coil_diameter_turns_100_200mm_step20_v4.csv")
    scan_200_300 = pd.read_csv(PRE / "coil_diameter_turns_200_300mm_step20_v4.csv")
    baseline = pd.read_csv(PRE / "baseline_220mm_10turns.csv")
    pitch_180_220 = pd.read_csv(PRE / "hex_pitch_opt_220mm_10turns_180_220" / "pitch_summary.csv")
    pitch_50_150 = pd.read_csv(PRE / "hex_pitch_opt_220mm_10turns_pitch50_150" / "pitch_summary.csv")

    best_100_200 = scan_100_200.sort_values("eta_theory", ascending=False).iloc[0]
    best_200_300 = scan_200_300.sort_values("eta_theory", ascending=False).iloc[0]
    best_pitch_overlap = pitch_50_150.sort_values("score", ascending=False).iloc[0]
    best_pitch_no_overlap = pitch_180_220.sort_values("score", ascending=False).iloc[0]
    base = baseline.iloc[0]

    doc = Document()
    doc.add_heading("AGV无线电能传输六边形阵列式耦合机构参数设计报告", level=0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("1. 论文目的与设计流程", level=1)
    doc.add_paragraph("本报告统一说明从论文目标出发，逐步确定固定参数、筛选线圈直径与匝数，再到六边形阵列间距优化的完整过程。")
    doc.add_paragraph("流程：目标定义 -> 固定参数确定 -> 变量范围扫描 -> 基础参数定型 -> 阵列间距优化 -> 工程约束修正。")

    doc.add_heading("2. 固定参数（设计起点）", level=1)
    add_table_from_df(doc, fixed_df, "固定参数表")

    doc.add_heading("3. 线圈直径与匝数筛选过程", level=1)
    doc.add_paragraph("第一阶段扫描范围：100~200 mm（步长20 mm），匝数10~20（步长2）。")
    doc.add_paragraph("第二阶段扫描范围：200~300 mm（步长20 mm），匝数10~20（步长2）。")
    add_table_from_df(
        doc,
        pd.DataFrame(
            [
                ["100~200 mm扫描最优", best_100_200["diameter_cm"] * 10, int(best_100_200["turns"]), best_100_200["eta_theory"], best_100_200["k"], best_100_200["L_air_uH"]],
                ["200~300 mm扫描最优", best_200_300["diameter_cm"] * 10, int(best_200_300["turns"]), best_200_300["eta_theory"], best_200_300["k"], best_200_300["L_air_uH"]],
            ],
            columns=["阶段", "直径(mm)", "匝数", "eta_theory", "k", "L_air(uH)"],
        ),
        "两阶段扫描结果对比",
        max_rows=10,
    )

    doc.add_paragraph(
        "结合工程约束（电感区间、交流损耗、尺寸可制造性）后，选定基础参数为：直径220 mm、10匝。"
    )
    add_table_from_df(doc, baseline, "基础参数点(220mm,10匝)计算结果", max_rows=5)

    doc.add_heading("4. 六边形阵列间距优化（偏移0~150 mm）", level=1)
    doc.add_paragraph("在固定直径220 mm、10匝条件下，针对偏移量0~150 mm进行阵列间距Pitch优化。")
    doc.add_paragraph("(A) 物理不考虑重叠约束，Pitch扫描50~150 mm。")
    add_table_from_df(doc, pitch_50_150, "Pitch=50~150 mm优化汇总", max_rows=20)
    doc.add_paragraph(
        f"该范围内评分最优Pitch为 {best_pitch_overlap['pitch_mm']:.0f} mm，"
        f"eta_min={best_pitch_overlap['eta_min']:.6f}，eta_mean={best_pitch_overlap['eta_mean']:.6f}。"
    )

    doc.add_paragraph("(B) 工程可制造约束（避免与220 mm直径线圈重叠），Pitch扫描180~220 mm。")
    add_table_from_df(doc, pitch_180_220, "Pitch=180~220 mm优化汇总", max_rows=20)
    doc.add_paragraph(
        f"在无重叠约束下，推荐Pitch为 {best_pitch_no_overlap['pitch_mm']:.0f} mm，"
        f"eta_min={best_pitch_no_overlap['eta_min']:.6f}，eta_mean={best_pitch_no_overlap['eta_mean']:.6f}。"
    )

    doc.add_heading("5. 最终建议参数", level=1)
    final_df = pd.DataFrame(
        [
            ["工作频率", "85 kHz"],
            ["输入电流", "20 A"],
            ["传输间隙", "50 mm"],
            ["线圈直径", "220 mm"],
            ["匝数", "10"],
            ["阵列形式", "六边形耦合思路下的1x3阵列"],
            ["Pitch(理论电磁最优)", f"{best_pitch_overlap['pitch_mm']:.0f} mm"],
            ["Pitch(工程推荐，无重叠)", f"{best_pitch_no_overlap['pitch_mm']:.0f} mm"],
        ],
        columns=["项目", "推荐值"],
    )
    add_table_from_df(doc, final_df, "最终参数建议表")

    doc.add_heading("6. 可追溯结果文件", level=1)
    doc.add_paragraph("以下文件已用于本报告：")
    refs = [
        PRE / "coil_diameter_turns_100_200mm_step20_v4.csv",
        PRE / "coil_diameter_turns_200_300mm_step20_v4.csv",
        PRE / "baseline_220mm_10turns.csv",
        PRE / "hex_pitch_opt_220mm_10turns_pitch50_150" / "pitch_summary.csv",
        PRE / "hex_pitch_opt_220mm_10turns_180_220" / "pitch_summary.csv",
        ROOT / "hex_array_redesign_blueprint_220mm_10turns.yaml",
    ]
    for p in refs:
        doc.add_paragraph(str(p))

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    print(f"Saved: {out_docx}")


if __name__ == "__main__":
    main()
