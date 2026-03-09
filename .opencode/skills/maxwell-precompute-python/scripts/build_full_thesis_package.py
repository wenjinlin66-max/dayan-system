from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Inches


def ensure_dirs(root: Path) -> dict[str, Path]:
    dirs = {
        "root": root,
        "manifest": root / "00_manifest",
        "report": root / "01_report",
        "fig_precompute": root / "02_figures" / "02_precompute",
        "fig_pitch": root / "02_figures" / "02_pitch_opt",
        "tables": root / "03_tables",
        "simulation": root / "04_simulation",
        "scripts": root / "06_scripts",
    }
    for p in dirs.values():
        p.mkdir(parents=True, exist_ok=True)
    return dirs


def copy_if_exists(src: Path, dst: Path) -> bool:
    if not src.exists():
        return False
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    return True


def add_table(doc: Any, df: pd.DataFrame, title: str, max_rows: int = 20) -> None:
    doc.add_paragraph(title)
    show = df.head(max_rows)
    if show.empty:
        doc.add_paragraph("(无数据)")
        return
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    for i, c in enumerate(show.columns):
        table.rows[0].cells[i].text = str(c)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(show.columns):
            cells[i].text = str(row[c])


def add_picture_safe(doc: Any, img: Path, caption: str, width: float = 5.8) -> None:
    if not img.exists():
        doc.add_paragraph(f"[缺失图片] {img}")
        return
    doc.add_paragraph(caption)
    doc.add_picture(str(img), width=Inches(width))


def heading(doc: Any, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def main() -> None:
    workspace = Path("results")
    agv_root = workspace / "agv_pdf_run"
    pre = agv_root / "precompute_outputs"

    date_tag = datetime.now().strftime("%Y%m%d")
    pkg_root = workspace / "final_thesis_package" / f"agv_wpt_thesis_{date_tag}"
    dirs = ensure_dirs(pkg_root)

    # Input data
    p_scan_100_200 = pre / "coil_diameter_turns_100_200mm_step20_v4.csv"
    p_scan_200_300 = pre / "coil_diameter_turns_200_300mm_step20_v4.csv"
    p_baseline = pre / "baseline_220mm_10turns.csv"
    p_pitch_50_150 = pre / "hex_pitch_opt_220mm_10turns_pitch50_150" / "pitch_summary.csv"
    p_pitch_180_220 = pre / "hex_pitch_opt_220mm_10turns_180_220" / "pitch_summary.csv"

    p_validate_status = agv_root / "validate_redesign_220mm_10turns" / "validate_autofix_status.json"
    p_autobuild_status = agv_root / "redesign_auto_build_220mm_10turns_v2" / "auto_build_status.json"
    p_retry_status = agv_root / "maxwell_outputs" / "quick_retry_20260226_0200" / "maxwell_run_status.json"

    scan_100_200 = pd.read_csv(p_scan_100_200)
    scan_200_300 = pd.read_csv(p_scan_200_300)
    baseline = pd.read_csv(p_baseline)
    pitch_50_150 = pd.read_csv(p_pitch_50_150)
    pitch_180_220 = pd.read_csv(p_pitch_180_220)

    with p_validate_status.open("r", encoding="utf-8") as f:
        validate_status = json.load(f)
    with p_autobuild_status.open("r", encoding="utf-8") as f:
        autobuild_status = json.load(f)
    with p_retry_status.open("r", encoding="utf-8") as f:
        retry_status = json.load(f)

    best_100_200 = scan_100_200.sort_values("eta_theory", ascending=False).iloc[0]
    best_200_300 = scan_200_300.sort_values("eta_theory", ascending=False).iloc[0]
    best_pitch_overlap = pitch_50_150.sort_values("score", ascending=False).iloc[0]
    best_pitch_no_overlap = pitch_180_220.sort_values("score", ascending=False).iloc[0]
    base = baseline.iloc[0]

    # Copy core assets
    for src in [p_scan_100_200, p_scan_200_300, p_baseline, p_pitch_50_150, p_pitch_180_220]:
        copy_if_exists(src, dirs["tables"] / src.name)

    for src in [p_validate_status, p_autobuild_status, p_retry_status, agv_root / "hex_array_redesign_blueprint_220mm_10turns.yaml"]:
        copy_if_exists(src, dirs["simulation"] / src.name)

    pre_figs = [
        pre / "coil_diameter_turns_100_200mm_step20_v4_efficiency_3d_cn.png",
        pre / "coil_diameter_turns_200_300mm_step20_v4_efficiency_3d_cn.png",
        pre / "baseline_220mm_10turns_efficiency_3d_cn.png",
    ]
    for fp in pre_figs:
        copy_if_exists(fp, dirs["fig_precompute"] / fp.name)

    pitch_fig_pairs = [
        (pre / "hex_pitch_opt_220mm_10turns_pitch50_150" / "fig_pitch_summary.png", dirs["fig_pitch"] / "fig_pitch_summary_pitch50_150.png"),
        (pre / "hex_pitch_opt_220mm_10turns_pitch50_150" / "fig_pitch_offset_heatmap.png", dirs["fig_pitch"] / "fig_pitch_offset_heatmap_pitch50_150.png"),
        (pre / "hex_pitch_opt_220mm_10turns_180_220" / "fig_pitch_summary.png", dirs["fig_pitch"] / "fig_pitch_summary_pitch180_220.png"),
        (pre / "hex_pitch_opt_220mm_10turns_180_220" / "fig_pitch_offset_heatmap.png", dirs["fig_pitch"] / "fig_pitch_offset_heatmap_pitch180_220.png"),
    ]
    for src, dst in pitch_fig_pairs:
        copy_if_exists(src, dst)

    copy_if_exists(Path(".opencode/skills/maxwell-precompute-python/scripts/diameter_turns_sweep.py"), dirs["scripts"] / "diameter_turns_sweep.py")
    copy_if_exists(Path(".opencode/skills/maxwell-precompute-python/scripts/hex_array_pitch_optimizer.py"), dirs["scripts"] / "hex_array_pitch_optimizer.py")
    copy_if_exists(Path(".opencode/skills/maxwell-precompute-python/scripts/build_full_thesis_package.py"), dirs["scripts"] / "build_full_thesis_package.py")

    # Build report (aligned to user-provided directory)
    doc = Document()
    heading(doc, "AGV无线电能传输阵列式耦合机构设计论文（目录化版本）", 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    heading(doc, "论文目录（建议）", 1)
    toc_lines = [
        "第一章 绪论",
        "第二章 无线电能传输系统理论基础",
        "第三章 耦合机构关键参数优化与阵列设计",
        "第四章 耦合机构的有限元仿真建模（Maxwell）",
        "第五章 结果分析与对比验证",
        "第六章 结论与展望",
    ]
    for t in toc_lines:
        doc.add_paragraph(t)

    # 第一章
    heading(doc, "第一章 绪论", 1)
    heading(doc, "1.1 研究背景及意义", 2)
    heading(doc, "1.1.1 智能物流与AGV的发展现状", 3)
    doc.add_paragraph("智能物流场景对AGV连续运行能力提出更高要求，传统接触式充电在维护成本、对位精度和可靠性方面存在限制。")
    heading(doc, "1.1.2 AGV无线充电技术的应用优势", 3)
    doc.add_paragraph("无线电能传输可降低机械接触磨损并支持一定偏移容忍，适用于多工位、频繁启停的AGV作业环境。")

    heading(doc, "1.2 无线电能传输技术综述", 2)
    heading(doc, "1.2.1 磁耦合谐振式技术原理", 3)
    doc.add_paragraph("磁耦合谐振系统通过发射与接收回路在工作频率附近的谐振增强能量耦合，核心指标为互感M与耦合系数k。")
    heading(doc, "1.2.2 阵列式耦合机构的研究现状", 3)
    doc.add_paragraph("阵列式结构通过空间冗余和激励切换提升偏移工况稳定性，但需在尺寸、损耗与控制复杂度之间做权衡。")

    heading(doc, "1.3 本文主要研究内容与技术路线", 2)
    heading(doc, "1.3.1 研究目标与关键问题", 3)
    doc.add_paragraph("目标是确定可落地的线圈直径、匝数与阵列间距参数，兼顾效率趋势与工程可制造性。")
    heading(doc, "1.3.2 理论计算与仿真结合的技术路线", 3)
    doc.add_paragraph("先通过Python参数化计算完成趋势筛选，再将候选方案导入Maxwell进行可建模与可求解验证。")

    # 第二章
    heading(doc, "第二章 无线电能传输系统理论基础", 1)
    heading(doc, "2.1 磁耦合谐振系统等效电路分析", 2)
    heading(doc, "2.1.1 SS型补偿拓扑结构推导", 3)
    doc.add_paragraph("本工作采用SS型补偿场景下的效率代理表达式进行参数比较，用于相同频率与负载条件下的相对优选。")
    heading(doc, "2.1.2 互感、耦合系数与传输效率的关系", 3)
    doc.add_paragraph("随着偏移增大，互感M通常下降，k=M/sqrt(L1L2)随之减小，效率代理值呈下降趋势。")

    heading(doc, "2.2 线圈磁场的数学建模", 2)
    heading(doc, "2.2.1 毕奥-萨伐尔定律在矩形/六边形线圈中的应用", 3)
    doc.add_paragraph("在参数筛选阶段，使用简化轴向磁场与等效互感表达式表征几何参数变化的影响。")
    heading(doc, "2.2.2 诺伊曼公式（Neumann Formula）计算互感原理", 3)
    doc.add_paragraph("互感通过几何相对位置决定，本文采用可计算近似式在不同偏移和间距下快速评估M的变化。")

    heading(doc, "2.3 基于Python的参数化理论计算平台搭建", 2)
    heading(doc, "2.3.1 计算流程与算法实现", 3)
    doc.add_paragraph("核心计算式：L = mu0*N^2*r*(ln(8r/a)-2)；M = mu0*pi*Ntx*Nrx*r^4/(2*(r^2+g^2)^(3/2))；k=M/sqrt(L1L2)；eta=((wM)^2*Rload)/((R1*(R2+Rload)+(wM)^2)*(R2+Rload))。")
    heading(doc, "2.3.2 理想环境下参数变化趋势分析（3D效率曲面图分析）", 3)
    add_picture_safe(doc, dirs["fig_precompute"] / "coil_diameter_turns_100_200mm_step20_v4_efficiency_3d_cn.png", "图2-1 100~200mm范围直径-匝数-效率三维图")
    add_picture_safe(doc, dirs["fig_precompute"] / "coil_diameter_turns_200_300mm_step20_v4_efficiency_3d_cn.png", "图2-2 200~300mm范围直径-匝数-效率三维图")

    heading(doc, "2.4 本章小结", 2)
    doc.add_paragraph("本章建立了参数筛选所需的理论模型与计算流程，为第三章参数优化奠定基础。")

    # 第三章
    heading(doc, "第三章 耦合机构关键参数优化与阵列设计", 1)
    heading(doc, "3.1 基础参数的选型与约束", 2)
    heading(doc, "3.1.1 运行频率(85kHz)与功率需求的确定", 3)
    doc.add_paragraph("固定工况采用85kHz、输入电流20A、负载10ohm，符合当前AGV场景下的设计边界。")
    heading(doc, "3.1.2 物理尺寸限制与气隙距离设定", 3)
    doc.add_paragraph("采用50mm气隙作为基线分析条件，参数筛选后再叠加结构可制造性约束。")

    heading(doc, "3.2 单体线圈参数的理论筛选（基于Python）", 2)
    heading(doc, "3.2.1 线圈直径与电感量的匹配计算", 3)
    compare_df = pd.DataFrame(
        [
            {
                "阶段": "100~200mm扫描最优",
                "直径(mm)": best_100_200["diameter_cm"] * 10,
                "匝数": int(best_100_200["turns"]),
                "eta_theory": best_100_200["eta_theory"],
                "k": best_100_200["k"],
                "L_air(uH)": best_100_200["L_air_uH"],
            },
            {
                "阶段": "200~300mm扫描最优",
                "直径(mm)": best_200_300["diameter_cm"] * 10,
                "匝数": int(best_200_300["turns"]),
                "eta_theory": best_200_300["eta_theory"],
                "k": best_200_300["k"],
                "L_air(uH)": best_200_300["L_air_uH"],
            },
            {
                "阶段": "工程定型",
                "直径(mm)": 220,
                "匝数": 10,
                "eta_theory": base["eta_theory"],
                "k": base["k"],
                "L_air(uH)": base["L_air_uH"],
            },
        ]
    )
    add_table(doc, compare_df, "表3-1 线圈直径与匝数筛选结果", max_rows=10)
    heading(doc, "3.2.2 匝数对耦合强度影响的量化分析", 3)
    add_picture_safe(doc, dirs["fig_precompute"] / "baseline_220mm_10turns_efficiency_3d_cn.png", "图3-1 基线点(220mm,10匝)效率示意")

    heading(doc, "3.3 六边形阵列拓扑结构设计", 2)
    heading(doc, "3.3.1 阵列布局的数学坐标建模", 3)
    doc.add_paragraph("采用1x3阵列等效坐标建模，发射单元中心按pitch均匀排布，接收端在0~150mm偏移区间内扫描。")
    heading(doc, "3.3.2 基于互感波动率的阵列间距（Pitch）优化", 3)
    doc.add_paragraph("定义M_eff=max(M_i)，并统计eta_min、eta_mean、coverage，评分函数为score=0.6*eta_min+0.3*eta_mean+0.1*coverage。")
    add_table(doc, pitch_50_150, "表3-2 Pitch=50~150mm优化结果（理论最优包含重叠）", max_rows=12)
    add_table(doc, pitch_180_220, "表3-3 Pitch=180~220mm优化结果（工程无重叠）", max_rows=12)
    add_picture_safe(doc, dirs["fig_pitch"] / "fig_pitch_summary_pitch50_150.png", "图3-2 Pitch鲁棒性曲线（50~150mm）")
    add_picture_safe(doc, dirs["fig_pitch"] / "fig_pitch_summary_pitch180_220.png", "图3-3 Pitch鲁棒性曲线（180~220mm）")

    heading(doc, "3.4 设计方案的最终定型", 2)
    final_df = pd.DataFrame(
        [
            {"项目": "线圈直径", "值": "220 mm"},
            {"项目": "匝数", "值": "10"},
            {"项目": "理论电磁最优Pitch", "值": f"{best_pitch_overlap['pitch_mm']:.0f} mm（存在重叠）"},
            {"项目": "工程推荐Pitch", "值": f"{best_pitch_no_overlap['pitch_mm']:.0f} mm（无重叠）"},
        ]
    )
    add_table(doc, final_df, "表3-4 最终定型参数")

    heading(doc, "3.5 本章小结", 2)
    doc.add_paragraph("第三章确定了基础线圈参数与阵列间距策略，给出了理论最优与工程可实现两套结果。")

    # 第四章
    heading(doc, "第四章 耦合机构的有限元仿真建模（Maxwell）", 1)
    heading(doc, "4.1 仿真模型的搭建", 2)
    heading(doc, "4.1.1 阵列线圈、铁氧体磁芯与屏蔽铝板的参数化建模", 3)
    doc.add_paragraph("当前自动化流程已完成阵列线圈与空气域等基础对象流程尝试，磁芯与屏蔽层将在后续稳定建模链路中补充。")
    heading(doc, "4.1.2 材料属性设置与网格剖分策略", 3)
    doc.add_paragraph("自动化日志显示材料赋值动作可执行（copper/air），但网格操作在当前环境中受gRPC接口异常影响，尚未稳定完成。")

    heading(doc, "4.2 求解设置与激励加载", 2)
    heading(doc, "4.2.1 涡流场（Eddy Current）求解器配置", 3)
    doc.add_paragraph("验证日志存在Eddy设置提醒，需在后续模型收敛前重新核查设置一致性。")
    heading(doc, "4.2.2 绞线（Stranded）激励加载及边界条件设定", 3)
    doc.add_paragraph("日志记录显示Current激励与Radiation边界在现阶段未稳定创建，导致验证未通过。")

    heading(doc, "4.3 典型工况下的仿真任务设定", 2)
    heading(doc, "4.3.1 正对工况性能分析", 3)
    doc.add_paragraph("当前quick retry任务未得到成功求解结果，处于模型/设置完备性修正阶段。")
    heading(doc, "4.3.2 多维度偏移工况的变量设定", 3)
    doc.add_paragraph("偏移工况参数已在理论计算中覆盖0~150mm，为后续有限元批处理提供了范围边界。")

    sim_df = pd.DataFrame(
        [
            {
                "文件": "validate_autofix_status.json",
                "ok": validate_status.get("ok"),
                "关键信息": validate_status.get("message", ""),
            },
            {
                "文件": "auto_build_status.json",
                "ok": autobuild_status.get("ok"),
                "关键信息": autobuild_status.get("message", ""),
            },
            {
                "文件": "quick_retry_20260226_0200/maxwell_run_status.json",
                "ok": retry_status.get("ok"),
                "关键信息": retry_status.get("message", ""),
            },
        ]
    )
    add_table(doc, sim_df, "表4-1 当前仿真验证状态（事实记录）", max_rows=10)

    heading(doc, "4.4 本章小结", 2)
    doc.add_paragraph("本章完成了建模流程证据汇总：当前自动化链路尚未形成成功求解，需先解决导电路径校验与gRPC保存问题。")

    # 第五章
    heading(doc, "第五章 结果分析与对比验证", 1)
    heading(doc, "5.1 理论计算与仿真结果的对比", 2)
    heading(doc, "5.1.1 电感、互感参数的修正与验证", 3)
    doc.add_paragraph("理论侧已形成L/M/k与效率代理的可重复计算结果；仿真侧当前未得到稳定导出参数，因此本节以趋势一致性与待验证项形式给出。")
    heading(doc, "5.1.2 磁芯与屏蔽层对系统性能的影响分析", 3)
    doc.add_paragraph("磁芯与屏蔽层作用机理可提升耦合并抑制漏磁，但需在仿真链路稳定后进行定量对比。")

    heading(doc, "5.2 阵列式机构抗偏移性能评价", 2)
    heading(doc, "5.2.1 水平偏移下效率稳定性曲线分析", 3)
    add_picture_safe(doc, dirs["fig_pitch"] / "fig_pitch_offset_heatmap_pitch50_150.png", "图5-1 50~150mm间距场景下偏移热力图")
    heading(doc, "5.2.2 垂直气隙波动下的系统鲁棒性分析", 3)
    doc.add_paragraph("本文基线气隙固定为50mm；气隙扰动鲁棒性作为下一阶段扩展项。")

    heading(doc, "5.3 磁场分布与电磁安全评估", 2)
    heading(doc, "5.3.1 磁感应强度云图分析（磁饱和检查）", 3)
    doc.add_paragraph("当前报告使用理论B_map指标进行先验评估，有限元云图待仿真链路修复后补齐。")
    heading(doc, "5.3.2 漏磁分布与安全标准校核", 3)
    doc.add_paragraph("漏磁与安全标准校核需要完整求解后处理结果支持，当前阶段仅保留方法路径。")

    heading(doc, "5.4 本章小结", 2)
    doc.add_paragraph("第五章完成了可得结果与待补验证项的边界划分，保证结论来源可追溯且不夸大。")

    # 第六章
    heading(doc, "第六章 结论与展望", 1)
    heading(doc, "6.1 全文总结", 2)
    doc.add_paragraph("本文完成了从目标定义、参数筛选到阵列间距优化的全流程设计，确定了220mm/10匝基础方案及工程推荐Pitch=180mm。")
    heading(doc, "6.2 创新点归纳", 2)
    doc.add_paragraph("创新点在于将理论参数扫描与阵列鲁棒性评分统一到Python流程，并形成可复现实验数据包与文档自动生成链。")
    heading(doc, "6.3 研究不足与后续改进方向", 2)
    doc.add_paragraph("后续重点是修复Maxwell自动化链路中的导电路径与gRPC保存问题，补齐全量工况有限元结果并完成理论-仿真闭环校准。")

    report_path = dirs["report"] / "AGV无线电能传输阵列耦合机构论文（六章目录版）.docx"
    doc.save(str(report_path))

    manifest = {
        "package_root": str(pkg_root),
        "report": str(report_path),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "selected_baseline": {"diameter_mm": 220, "turns": 10},
        "pitch_best_overlap": float(best_pitch_overlap["pitch_mm"]),
        "pitch_best_engineering": float(best_pitch_no_overlap["pitch_mm"]),
        "simulation_status": {
            "validate_ok": validate_status.get("ok"),
            "autobuild_ok": autobuild_status.get("ok"),
            "retry_ok": retry_status.get("ok"),
        },
    }
    with (dirs["manifest"] / "package_manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)

    trace_text = (
        "# Source Traceability\n\n"
        "Core formulas and optimization methods are from:\n"
        "- .opencode/skills/maxwell-precompute-python/scripts/diameter_turns_sweep.py\n"
        "- .opencode/skills/maxwell-precompute-python/scripts/hex_array_pitch_optimizer.py\n"
        "Simulation evidence files:\n"
        "- results/agv_pdf_run/validate_redesign_220mm_10turns/validate_autofix_status.json\n"
        "- results/agv_pdf_run/redesign_auto_build_220mm_10turns_v2/auto_build_status.json\n"
        "- results/agv_pdf_run/maxwell_outputs/quick_retry_20260226_0200/maxwell_run_status.json\n"
    )
    (dirs["manifest"] / "source_traceability.md").write_text(trace_text, encoding="utf-8")

    print(f"Saved: {report_path}")
    print(f"Saved: {dirs['manifest'] / 'package_manifest.json'}")
    print(f"Saved: {dirs['manifest'] / 'source_traceability.md'}")
    print(f"Package root: {pkg_root}")


if __name__ == "__main__":
    main()
